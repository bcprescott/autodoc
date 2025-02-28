import os
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from openai import AzureOpenAI

load_dotenv()

def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from a DOCX file and return as a single string.

    :param docx_path: Path to the DOCX file
    :return: A single string containing the text from the document
    """
    document = Document(docx_path)
    return "\n".join(para.text.strip() for para in document.paragraphs if para.text.strip())

def create_formatted_docx(template_text: str, output_path: str):
    """
    Creates a formatted DOCX file from a templated text.

    :param template_text: A string containing the templated text
    :param output_path: The path to which the formatted document will be saved
    """
    
    doc = Document()
    lines = template_text.splitlines()
    content_buffer = []

    def flush_content():
        nonlocal content_buffer
        if content_buffer:
            content = "\n\n".join(content_buffer).strip()
            if content:
                doc.add_paragraph(content)
            content_buffer = []

    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue
        if stripped_line.endswith('**'):
            flush_content()
            header_text = stripped_line.rstrip('**').strip()
            heading = doc.add_heading(header_text, level=2)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            content_buffer.append(stripped_line)

    flush_content()
    doc.save(output_path)

def generate_template_from_document(doc_text: str, claim_data) -> str:
    """
    Generate a template version of a given document text using an LLM.

    Given a document text and some claim data, this function generates a template version of the document using an LLM.
    The LLM is given a prompt that asks it to identify the main sections of the document and replace specific details with placeholders.
    The LLM is also given some additional instructions to follow, such as adding new lines/spaces between each section/paragraph and keeping the headers and overall structure intact.

    :param doc_text: The text of the document to be templated
    :param claim_data: A dictionary containing some example claim data to use as placeholder values
    :return: A string containing the templated version of the document
    """
    prompt = (
        "Below is the content of an insurance claim letter document:\n\n"
        f"{doc_text}\n\n"
        "Please generate a template version of this document. Identify the main sections and replace specific details with placeholders."
        "Keep in mind that this document is addressed to the Claimant from the Insurance company."
        f"""Follow the additional instructions below:
        1. Add new lines/spaces between each section/paragraph.
        2. Add two '**' symbols after any section headers, such as 'Authorization Letter to Claim**'.
        3. Keep the headers and overall structure intact.    
        4. Do not respond with any extra text/information outside of the document template.
        5. The claimant is called "Claimant", the person writing the letter is called "Sender", the insurance company is called "Insurance Company".
        6. The letter begins with the claimant's name and address, followed by the insurance company's name and address.    
        7. Use the following claim example for placeholder names/schema: {claim_data}
        """
        "Replace the template sections with double curly braces. For example, if the document includes a line like 'Claim Number: 123456', replace it with 'Claim Number: {{ Insert Claim Number }}'."
    )
    
    response = client.chat.completions.create(
        model=deployment_id,
        messages=[
            {"role": "system", "content": "You are a helpful insurance claims assistant, helping convert example claims letters into reusable templates."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=800,
        temperature=0
    )
    
    return response.choices[0].message.content

if __name__ == "__main__":
    # Load environment variables for your Azure OpenAI specifics
    endpoint = os.getenv("azopenai_endpoint")
    api_key = os.getenv("azopenai_key")
    deployment_id = os.getenv("azopenai_deployment")
    api_version = os.getenv("azopenai_api_version")

    if not endpoint or not api_key or not deployment_id:
        raise EnvironmentError("Please set azopenai_endpoint, azopenai_key, and azopenai_deployment.")

    client = AzureOpenAI(api_key=api_key, api_version=api_version, azure_endpoint=endpoint)


    # Load the claim data from a JSON file
    with open("claim_extract.json", 'r') as file:
        claim_data = json.load(file)

    # Extract text from the example DOCX file (treating as a historic example)
    docx_file_path = "synthetic_letter_example.docx"
    document_text = extract_text_from_docx(docx_file_path)

    # Generate the template version of the document
    template_version = generate_template_from_document(document_text, claim_data)
    
    # Create a formatted DOCX file from the generated template version
    output_file = "formatted_letter_template.docx"
    create_formatted_docx(template_version, output_file)
