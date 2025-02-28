import json
import re
import os
from docx import Document
from openai import AzureOpenAI
from dotenv import load_dotenv

load_dotenv()

def replace_placeholders_in_paragraph(paragraph, replacements):
    """
    Replace placeholders in a paragraph's runs.
    Note: This simple approach assumes the placeholder is not split across multiple runs.
    """
    for key, value in replacements.items():
        pattern = r"{{\s*" + re.escape(key) + r"\s*}}"
        for run in paragraph.runs:
            run.text = re.sub(pattern, value, run.text)

def replace_placeholders_in_document(doc, replacements):
    """
    Iterate over all paragraphs (and those within tables) in the document and replace placeholders.
    """
    for para in doc.paragraphs:
        replace_placeholders_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_placeholders_in_paragraph(para, replacements)

def generate_injury_description(claim_text):
    prompt = (
        f"""
        Below is the content of an insurance claim from the claimant:
        
        <claim_information>
        {claim_text}
        </claim_information>
        
        Please generate a brief description of the injury/condition that occurred. Do not personalize the description.
        """
    )
    response = client.chat.completions.create(
        model=deployment_id,
        messages=[
            {"role": "system", "content": "You are a helpful insurance claims assistant, helping to analyze claim information and writing a detailed description of the injury."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=800,
        temperature=0
    )
    return response.choices[0].message.content

if __name__ == "__main__":

    # Load environment variables for your Azure OpenAI specifics
    endpoint = os.getenv("azopenai_endpoint")
    api_key = os.getenv("azopenai_key")
    deployment_id = os.getenv("azopenai_deployment")
    api_version = os.getenv("azopenai_api_version")
    client = AzureOpenAI(api_key=api_key, api_version=api_version, azure_endpoint=endpoint)

    # Load the claim information and generate the injury description
    with open("claim_extract.json") as f:
        replacements = json.load(f)

    # Generate the injury description
    injury_description = generate_injury_description(replacements)
    
    # Add the injury description to the replacements dictionary
    replacements["Description of Injury/Condition"] = injury_description

    # Load the template document and replace placeholders
    doc = Document("formatted_letter_template.docx")

    # Replace placeholders
    replace_placeholders_in_document(doc, replacements)

    # Save the populated document
    output_path = "populated_claim_letter.docx"
    doc.save(output_path)
    print(f"Populated document saved as {output_path}")
